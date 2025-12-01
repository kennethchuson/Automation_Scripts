import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext, ttk
import os
import threading
import re
import threading

# Resume reading
from docx import Document as DocxDocument
import PyPDF2

# LangChain + Ollama
from langchain_community.llms import Ollama
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain

# Output formats
from docx import Document as OutDocx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle


# ==================== RESUME TEXT EXTRACTION ====================
def extract_text_from_file(path):
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        doc = DocxDocument(path)
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    elif ext == ".pdf":
        text = ""
        with open(path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text
    elif ext == ".txt":
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    else:
        raise ValueError("Only .docx, .pdf, .txt supported")


def detect_company_name(job_description: str) -> str:
    prompt = f"""
Extract ONLY the company name from this job description. 
Look for phrases like "at Google", "join Microsoft", "Company XYZ is hiring", "About Acme Inc", etc.
Return ONLY the company name, nothing else. If unsure, return "Company".

Job Description:
{job_description[:3000]}  # first part is enough
"""

    try:
        llm = Ollama(model="gemma3:4b", temperature=0.1)
        chain = LLMChain(llm=llm, prompt=PromptTemplate.from_template(prompt))
        result = chain.run({})
        name = result.strip().strip('"').strip("'").strip()
        if name.lower() in ["company", "hiring", "we", "our team", ""]:
            return "Company"
        return name
    except:
        return "Company"


def generate_cover_letter(resume_text: str, job_description: str) -> str:
    template = """
You are a world-class career writer.

Rules:
- Use ONLY real details from the resume: Full Name, Email, Phone, Location
- NEVER invent or guess anything
- NEVER use placeholders
- If name not found → start with "Dear Hiring Manager,"

Write a powerful, concise (320–420 words), ATS-friendly cover letter that perfectly matches the candidate to the role.

=== RESUME ===
{resume}

=== JOB DESCRIPTION ===
{job}
"""

    prompt = PromptTemplate.from_template(template)
    llm = Ollama(model="gemma3:4b", temperature=0.35)
    chain = LLMChain(llm=llm, prompt=prompt)
    return chain.run(resume=resume_text, job=job_description)


# ==================== SAVE DOCX ====================
def save_as_docx(text: str, filepath: str):
    doc = OutDocx()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        p = doc.add_paragraph(line)
        if "@" in line or re.match(r"[\d\s\-\(\)]{8,}", line) or len(line) < 60:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].bold = True
            p.space_after = Pt(12)
        else:
            p.space_after = Pt(8)
    doc.save(filepath)


# ==================== SAVE PDF ====================
def save_as_pdf(text: str, filepath: str):
    doc = SimpleDocTemplate(filepath, pagesize=letter,
                            rightMargin=72, leftMargin=72,
                            topMargin=80, bottomMargin=72)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='CenterBold',
                             parent=styles['Normal'],
                             alignment=1,
                             fontSize=11,
                             spaceAfter=12,
                             fontName='Helvetica-Bold'))

    story = []
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            story.append(Spacer(1, 6))
            continue
        if "@" in line or re.match(r"[\d\s\-\(\)]{8,}", line) or len(line) < 60:
            story.append(Paragraph(line, styles['CenterBold']))
        else:
            story.append(Paragraph(line, styles['Normal']))
        story.append(Spacer(1, 6))
    doc.build(story)


class CoverLetterApp:
    def __init__(self, master):
        self.master = master
        master.title("Auto Cover Letter Generator — gemma3:4b")
        master.geometry("920x760")
        master.resizable(True, True)

        tk.Label(master, text="1. Select Your Resume", font=("Arial", 12, "bold")).pack(anchor="w", padx=25, pady=(25,8))
        rframe = tk.Frame(master)
        rframe.pack(fill="x", padx=25, pady=5)
        self.resume_var = tk.StringVar()
        tk.Entry(rframe, textvariable=self.resume_var, width=80).pack(side="left", fill="x", expand=True, padx=(0,10))
        tk.Button(rframe, text="Browse Resume", command=self.browse).pack(side="right")

        tk.Label(master, text="2. Paste Full Job Description", font=("Arial", 12, "bold")).pack(anchor="w", padx=25, pady=(20,8))
        self.job_box = scrolledtext.ScrolledText(master, height=14, wrap=tk.WORD, font=("Segoe UI", 10))
        self.job_box.pack(fill="both", expand=True, padx=25, pady=5)

        self.gen_btn = tk.Button(master, text="Generate Cover Letter (DOCX + PDF)", font=("Arial", 14, "bold"),
                                 bg="#006400", fg="white", height=2, command=self.start)
        self.gen_btn.pack(pady=30)

        self.status = tk.StringVar(value="Ready – gemma3:4b loaded")
        tk.Label(master, textvariable=self.status, fg="#006400", font=("Arial", 11)).pack(fill="x", padx=25)

        self.progress = ttk.Progressbar(master, mode='indeterminate')
        self.progress.pack(fill="x", padx=25, pady=10)

    def browse(self):
        path = filedialog.askopenfilename(filetypes=[("Documents", "*.docx *.pdf *.txt")])
        if path:
            self.resume_var.set(path)
            self.status.set(f"Resume loaded: {os.path.basename(path)}")

    def start(self):
        resume_path = self.resume_var.get().strip()
        job_desc = self.job_box.get("1.0", tk.END).strip()

        if not resume_path or not os.path.exists(resume_path):
            messagebox.showwarning("Error", "Please select a valid resume file")
            return
        if not job_desc:
            messagebox.showwarning("Error", "Please paste the job description")
            return

        self.gen_btn.config(state="disabled")
        self.progress.start()
        threading.Thread(target=self.worker, args=(resume_path, job_desc), daemon=True).start()

    def worker(self, resume_path, job_desc):
        try:
            self.master.after(0, lambda: self.status.set("Reading resume..."))
            resume_text = extract_text_from_file(resume_path)

            self.master.after(0, lambda: self.status.set("Detecting company name from job post..."))
            company_name = detect_company_name(job_desc)
            self.master.after(0, lambda: self.status.set(f"Company detected: {company_name}"))

            self.master.after(0, lambda: self.status.set("Generating cover letter with gemma3:4b..."))
            cover_letter = generate_cover_letter(resume_text, job_desc)

            resume_name = os.path.splitext(os.path.basename(resume_path))[0]
            safe_company = re.sub(r'[^\w\-_]', '_', company_name)
            base_name = f"cover_letter_{resume_name}_{safe_company}"

            docx_path = os.path.join(os.getcwd(), f"{base_name}.docx")
            pdf_path = os.path.join(os.getcwd(), f"{base_name}.pdf")

            self.master.after(0, lambda: self.status.set("Saving DOCX..."))
            save_as_docx(cover_letter, docx_path)

            self.master.after(0, lambda: self.status.set("Saving PDF..."))
            save_as_pdf(cover_letter, pdf_path)

            self.master.after(0, lambda: self.status.set(f"Done! Generated 2 files → {company_name}"))
            messagebox.showinfo("Success!",
                                f"Cover letter generated for {company_name}!\n\n"
                                f"• {os.path.basename(docx_path)}\n"
                                f"• {os.path.basename(pdf_path)}")

            folder = os.path.dirname(docx_path)
            os.startfile(folder) if os.name == "nt" else None

        except Exception as e:
            messagebox.showerror("Error", f"Something went wrong:\n{e}")
            self.master.after(0, lambda: self.status.set("Failed"))
        finally:
            self.master.after(0, lambda: self.gen_btn.config(state="normal"))
            self.master.after(0, lambda: self.progress.stop())


if __name__ == "__main__":
    root = tk.Tk()
    app = CoverLetterApp(root)
    root.mainloop()
